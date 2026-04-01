import json
import logging
import os
import re
import tempfile
from pathlib import Path
from typing import List, Optional, Any
from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
import structlog
import asyncio

from app.db.database import get_session
from app.services.agent_service import AgentService
from app.api.settings import get_active_llm_settings
from app.services.mcp_builder_service import mcp_builder_service

logger = structlog.get_logger()
router = APIRouter()


def _extract_supabase_jwt(request: Request) -> str | None:
    """Extract Supabase JWT from Authorization header for gateway auth.

    Starlette's BaseHTTPMiddleware runs endpoints in a separate asyncio task,
    so ContextVars set in the middleware don't propagate to non-streaming
    handlers.  We extract the token directly from the request instead.
    """
    auth = request.headers.get("Authorization") or ""
    if auth.startswith("Bearer "):
        return auth[7:].strip() or None
    return None


# ---------------------------------------------------------------------------
# In-memory generation progress store (survives across SSE disconnect)
# ---------------------------------------------------------------------------
_generation_progress: dict[str, dict] = {}
# Key: server_name (safe_product_name)
# Value: { "status": "running"|"completed"|"failed",
#           "phase": str, "logs": list[str], "iteration": int,
#           "result": dict|None }


@router.get("/progress/{server_name}")
async def get_generation_progress(server_name: str):
    """Get the current generation progress for a server build."""
    progress = _generation_progress.get(server_name)
    if not progress:
        return {"status": "not_found"}
    return progress


# ---------------------------------------------------------------------------
# Document Upload & Text Extraction
# ---------------------------------------------------------------------------

def _extract_pdf_smart(path: str) -> tuple[str, str]:
    """Per-page smart PDF extraction.

    - Pages with NO images → pymupdf text extraction (fast, accurate)
    - Pages WITH images → OCR the whole page (catches all visible text)

    Returns (extracted_text, method) where method is "text", "ocr", or "hybrid".
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise HTTPException(status_code=500, detail="pymupdf not installed. Cannot process PDFs.")

    doc = fitz.open(path)
    pages_text: list[str] = []
    ocr_page_count = 0
    text_page_count = 0

    try:
        for page_num, page in enumerate(doc):
            images = page.get_images(full=True)

            if not images:
                # No images on this page → fast pymupdf text extraction
                text = page.get_text() or ""
                if text.strip():
                    pages_text.append(text.strip())
                    text_page_count += 1
            else:
                # Page has images → OCR the whole page to catch everything
                ocr_text = _ocr_page(page)
                if ocr_text and ocr_text.strip():
                    pages_text.append(ocr_text.strip())
                    ocr_page_count += 1
                else:
                    # OCR failed/unavailable — fall back to pymupdf text
                    text = page.get_text() or ""
                    if text.strip():
                        pages_text.append(text.strip())
                        text_page_count += 1
    finally:
        doc.close()

    combined = "\n\n".join(pages_text)
    total = text_page_count + ocr_page_count

    if ocr_page_count == 0:
        method = "text"
    elif text_page_count == 0:
        method = "ocr"
    else:
        method = "hybrid"

    logger.info("pdf_smart_extract", pages=total, text_pages=text_page_count, ocr_pages=ocr_page_count, chars=len(combined), method=method)
    return combined, method


def _ocr_page(page) -> Optional[str]:
    """Render a pymupdf page to image and OCR it with Tesseract."""
    try:
        import pytesseract
        from PIL import Image
        import io
    except ImportError:
        logger.warning("ocr_deps_missing", hint="Install pytesseract and Pillow for OCR support")
        return None

    try:
        # Check if Tesseract is available
        pytesseract.get_tesseract_version()
    except Exception:
        logger.warning("tesseract_not_installed", hint="brew install tesseract")
        return None

    try:
        # Render page to image at 200 DPI (good balance of speed vs quality)
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))

        # Run OCR
        text = pytesseract.image_to_string(img)
        return text
    except Exception as exc:
        logger.warning("ocr_page_failed", error=str(exc))
        return None


def _extract_docx(path: str) -> str:
    """Extract text from DOCX including paragraphs and tables."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Cannot process DOCX files.")

    try:
        doc = DocxDocument(path)
        parts: list[str] = []

        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        # Extract tables (API docs often use tables for parameters)
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_rows.append(row_text)
            if table_rows:
                parts.append("\n".join(table_rows))

        return "\n\n".join(parts)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX file: {str(exc)}")


@router.post("/upload-docs")
async def upload_documentation(file: UploadFile = File(...)):
    """Upload a documentation file (PDF, DOCX, TXT, JSON, YAML) and extract text."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    ext = os.path.splitext(file.filename)[1].lower()
    allowed = {".pdf", ".docx", ".doc", ".txt", ".md", ".json", ".yaml", ".yml"}

    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Allowed: {', '.join(sorted(allowed))}")

    logger.info("mcp_builder_upload_start", filename=file.filename, ext=ext)

    # Save to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        pages = 0
        method = "text"

        if ext == ".pdf":
            text, method = _extract_pdf_smart(tmp_path)
            # Count pages
            try:
                import fitz
                doc = fitz.open(tmp_path)
                pages = len(doc)
                doc.close()
            except Exception:
                pass

        elif ext in (".docx", ".doc"):
            text = _extract_docx(tmp_path)
            pages = 1  # DOCX doesn't have meaningful page count for our purposes

        else:
            # Text files — read as UTF-8
            try:
                with open(tmp_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
            except Exception as exc:
                raise HTTPException(status_code=400, detail=f"Failed to read file: {str(exc)}")

        if not text or not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from this file. For scanned PDFs, ensure Tesseract is installed (brew install tesseract)."
            )

        logger.info("mcp_builder_upload_complete", filename=file.filename, chars=len(text), pages=pages, method=method)

        return {
            "text": text,
            "pages": pages,
            "method": method,
            "filename": file.filename,
        }

    finally:
        # Clean up temp file
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# MCP Server Registry
# ---------------------------------------------------------------------------

def _check_server_status(destination_folder: str, module_name: str) -> str:
    """Check if an MCP server is running, stopped, or has errors."""
    import subprocess

    server_path = os.path.expanduser(f"{destination_folder}/src/{module_name}/server.py")

    # Check if files exist
    if not os.path.exists(server_path):
        return "error"

    # Check if process is running
    try:
        result = subprocess.run(
            ["pgrep", "-f", f"{module_name}.server"],
            capture_output=True, text=True, timeout=5
        )
        if result.returncode == 0:
            return "running"
    except Exception:
        pass

    return "stopped"


@router.get("/servers")
async def list_mcp_servers(session: AsyncSession = Depends(get_session)):
    """List all built MCP servers with live status."""
    from app.models.mcp_server import McpServer
    from sqlalchemy import select

    result = await session.execute(
        select(McpServer).order_by(McpServer.created_at.desc())
    )
    servers = result.scalars().all()

    server_list = []
    for srv in servers:
        data = srv.to_dict()
        data["status"] = _check_server_status(srv.destination_folder, srv.module_name)
        server_list.append(data)

    return server_list


@router.get("/servers/{server_id}")
async def get_mcp_server(server_id: str, session: AsyncSession = Depends(get_session)):
    """Get a single MCP server's details with live status."""
    from app.models.mcp_server import McpServer
    from sqlalchemy import select

    result = await session.execute(
        select(McpServer).where(McpServer.server_id == server_id)
    )
    srv = result.scalar_one_or_none()
    if not srv:
        raise HTTPException(status_code=404, detail="Server not found")

    data = srv.to_dict()
    data["status"] = _check_server_status(srv.destination_folder, srv.module_name)
    return data


@router.delete("/servers/{server_id}")
async def delete_mcp_server(
    server_id: str,
    delete_files: bool = False,
    session: AsyncSession = Depends(get_session),
):
    """Remove an MCP server from the registry. Optionally delete generated files on disk."""
    import shutil
    from app.models.mcp_server import McpServer
    from sqlalchemy import select

    result = await session.execute(
        select(McpServer).where(McpServer.server_id == server_id)
    )
    srv = result.scalar_one_or_none()
    if not srv:
        raise HTTPException(status_code=404, detail="Server not found")

    files_deleted = False
    if delete_files and srv.destination_folder:
        folder = srv.destination_folder
        if folder.startswith("~/") or folder.startswith("~\\"):
            folder = str(Path.home() / folder[2:])
        folder_path = Path(folder).resolve()

        # Safety: only delete if the folder's parent is named "mcp-servers"
        # e.g. ~/mcp-servers/gta ✓, ~/Desktop/mcp-servers/gta ✓
        # e.g. ~/Desktop ✗, ~/mcp-servers ✗, / ✗
        is_safe = (folder_path.parent.name == "mcp-servers"
                   and len(folder_path.parts) >= 3
                   and folder_path.name != "mcp-servers")

        if is_safe and folder_path.is_dir():
            shutil.rmtree(folder_path, ignore_errors=True)
            files_deleted = True
        elif not is_safe and folder_path.is_dir():
            logger.warning("delete_mcp_server_unsafe_path", path=str(folder_path),
                           reason="Path parent is not 'mcp-servers/', refusing to delete")

    await session.delete(srv)
    return {"status": "deleted", "server_id": server_id, "files_deleted": files_deleted}


# ---------------------------------------------------------------------------
# MCP Server Start / Stop
# ---------------------------------------------------------------------------

_running_servers: dict[str, "subprocess.Popen[bytes]"] = {}


def _expand_dest(destination_folder: str) -> str:
    """Expand ~ in destination folder paths."""
    if destination_folder.startswith("~/") or destination_folder.startswith("~\\"):
        return str(Path.home() / destination_folder[2:])
    return os.path.expanduser(destination_folder)


@router.post("/servers/{server_id}/start")
async def start_mcp_server(server_id: str, session: AsyncSession = Depends(get_session)):
    """Start an MCP server process."""
    import subprocess
    from app.models.mcp_server import McpServer
    from sqlalchemy import select

    result = await session.execute(
        select(McpServer).where(McpServer.server_id == server_id)
    )
    srv = result.scalar_one_or_none()
    if not srv:
        raise HTTPException(status_code=404, detail="Server not found")

    dest = _expand_dest(srv.destination_folder)
    module_name = srv.module_name

    # Check if already running
    if _check_server_status(srv.destination_folder, module_name) == "running":
        return {"status": "already_running", "server_id": server_id}

    venv_python = os.path.join(dest, ".venv", "bin", "python")
    if os.path.isfile(venv_python):
        cmd = [venv_python, "-m", f"{module_name}.server"]
    else:
        # No venv yet — use uv run which auto-creates one
        cmd = ["uv", "run", "--directory", dest, "python", "-m", f"{module_name}.server"]

    try:
        proc = subprocess.Popen(
            cmd,
            cwd=dest,
            stdin=subprocess.PIPE,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            start_new_session=True,
        )
        _running_servers[server_id] = proc
        return {"status": "started", "server_id": server_id, "pid": proc.pid}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to start server: {e}")


@router.post("/servers/{server_id}/stop")
async def stop_mcp_server(server_id: str, session: AsyncSession = Depends(get_session)):
    """Stop a running MCP server process."""
    import subprocess
    import signal
    from app.models.mcp_server import McpServer
    from sqlalchemy import select

    result = await session.execute(
        select(McpServer).where(McpServer.server_id == server_id)
    )
    srv = result.scalar_one_or_none()
    if not srv:
        raise HTTPException(status_code=404, detail="Server not found")

    killed = False

    # Try killing from tracked processes first
    proc = _running_servers.pop(server_id, None)
    if proc and proc.poll() is None:
        try:
            os.killpg(os.getpgid(proc.pid), signal.SIGTERM)
            proc.wait(timeout=5)
            killed = True
        except Exception:
            try:
                proc.kill()
                killed = True
            except Exception:
                pass

    # Fallback: kill by pattern
    if not killed:
        module_name = srv.module_name
        try:
            subprocess.run(
                ["pkill", "-f", f"{module_name}.server"],
                capture_output=True, timeout=5,
            )
            killed = True
        except Exception:
            pass

    return {"status": "stopped", "server_id": server_id, "killed": killed}


class MCPDiscoverRequest(BaseModel):
    product_id: str

class MCPDiscoverExternalRequest(BaseModel):
    api_name: str
    api_docs_url: Optional[str] = None
    api_docs_text: Optional[str] = None

class MCPGeneratorRequest(BaseModel):
    product_id: Optional[str] = None # Native product ID or None for external
    product_name: str # Kept product_name as it's used in task_prompt and module_name
    destination_folder: str
    selected_endpoints: List[dict] # Full list of selected endpoints with schema
    api_docs_url: Optional[str] = None
    api_docs_text: Optional[str] = None # Added api_docs_text
    api_base_url: Optional[str] = None
    auth_type: Optional[str] = None
    auth_details: Optional[str] = None
    kb_product_id: Optional[str] = None

@router.post("/discover")
async def discover_mcp_endpoints(
    request: MCPDiscoverRequest,
    raw_request: Request,
    session: AsyncSession = Depends(get_session)
):
    logger.info("mcp_builder_discover_start", product_id=request.product_id)

    llm_settings = await get_active_llm_settings(session, supabase_jwt=_extract_supabase_jwt(raw_request))
    api_key = llm_settings.get("api_key")
    if not api_key:
        detail = llm_settings.get("llm_unavailable_detail", "LLM API key is required to analyze the knowledge base.")
        raise HTTPException(status_code=400, detail=detail)

    model = llm_settings.get("model_name", "gpt-4o")
    api_url = llm_settings.get("api_url")
    default_headers = llm_settings.get("default_headers", {})

    endpoints = await mcp_builder_service.discover_endpoints(
        product_id=request.product_id,
        api_key=api_key,
        model=model,
        api_url=api_url,
        default_headers=default_headers,
        provider=llm_settings.get("provider", "openai"),
    )
    
    return {"status": "success", "endpoints": [ep.model_dump() for ep in endpoints]}

@router.post("/discover-external")
async def discover_external_endpoints(
    request: MCPDiscoverExternalRequest,
    raw_request: Request,
    session: AsyncSession = Depends(get_session)
):
    """Discover endpoints from an external API docs URL."""
    logger.info("mcp_builder_discover_external_start", api_name=request.api_name, url=request.api_docs_url)

    llm_settings = await get_active_llm_settings(session, supabase_jwt=_extract_supabase_jwt(raw_request))
    api_key = llm_settings.get("api_key")
    if not api_key:
        detail = llm_settings.get("llm_unavailable_detail", "LLM API key is required to analyze the API documentation.")
        raise HTTPException(status_code=400, detail=detail)

    model = llm_settings.get("model_name", "gpt-4o")
    api_url = llm_settings.get("api_url")
    default_headers = llm_settings.get("default_headers", {})

    if not request.api_docs_url and not request.api_docs_text:
        raise HTTPException(status_code=400, detail="Must provide either api_docs_url or api_docs_text")

    if request.api_docs_text:
        result = await mcp_builder_service.discover_from_text(
            api_name=request.api_name,
            api_docs_text=request.api_docs_text,
            api_key=api_key,
            model=model,
            api_url=api_url,
            default_headers=default_headers
        )
    else:
        result = await mcp_builder_service.discover_from_url(
            api_name=request.api_name,
            api_docs_url=request.api_docs_url,
            api_key=api_key,
            model=model,
            api_url=api_url,
            default_headers=default_headers
        )

    # Determine kb_product_id from the result
    safe_name = re.sub(r'[^a-z0-9]', '_', request.api_name.lower()).strip('_')
    from app.services.mcp_builder_service import MCP_PRODUCT_PREFIX
    kb_product_id = f"{MCP_PRODUCT_PREFIX}{safe_name}" if result.raw_pages else None

    return {
        "status": "success",
        "base_url": result.base_url,
        "auth_type": result.auth_type,
        "auth_details": result.auth_details,
        "endpoints": [ep.model_dump() for ep in result.endpoints],
        "pages_crawled": result.pages_crawled,
        "kb_product_id": kb_product_id,
    }

class MCPDiscoverAgentRequest(BaseModel):
    api_name: str
    api_docs_url: Optional[str] = None
    api_docs_text: Optional[str] = None


@router.post("/discover-agent")
async def discover_agent_endpoints(
    request: MCPDiscoverAgentRequest,
    raw_request: Request,
    session: AsyncSession = Depends(get_session)
):
    """Fast crawl + KB storage: backend crawler fetches docs, stores in KB, returns kb_product_id."""
    logger.info("mcp_builder_fetch_docs_start", api_name=request.api_name, url=request.api_docs_url)

    if not request.api_docs_url and not request.api_docs_text:
        raise HTTPException(status_code=400, detail="Must provide either api_docs_url or api_docs_text")

    safe_name = re.sub(r'[^a-z0-9]', '_', request.api_name.lower()).strip('_')
    from app.services.mcp_builder_service import store_mcp_docs_kb, _crawl_api_docs, _extract_text_from_html

    pages: list[tuple[str, str]] = []

    if request.api_docs_text:
        # Text/PDF upload — already have the content
        pages.append(("Uploaded API Documentation", request.api_docs_text))
    elif request.api_docs_url:
        # URL — fast backend crawl
        import httpx as _httpx

        # Fetch initial page
        async with _httpx.AsyncClient(timeout=15.0, follow_redirects=True) as http:
            resp = await http.get(request.api_docs_url, headers={
                "User-Agent": "ReTrace-MCP-Builder/1.0",
                "Accept": "text/html, application/json, */*",
            })
            resp.raise_for_status()
            initial_html = resp.text
            content_type = resp.headers.get("content-type", "")

        # Check if it's an OpenAPI spec
        if "json" in content_type or request.api_docs_url.endswith((".json", ".yaml", ".yml")):
            pages.append((f"{request.api_name} - OpenAPI Spec", initial_html))
        elif "html" in content_type:
            # Check for embedded OpenAPI spec link
            import re as _re
            spec_match = _re.search(r'href="([^"]*(?:openapi|swagger)[^"]*\.(?:json|yaml|yml))"', initial_html, _re.IGNORECASE)
            if spec_match:
                spec_url = spec_match.group(1)
                if not spec_url.startswith("http"):
                    from urllib.parse import urljoin
                    spec_url = urljoin(request.api_docs_url, spec_url)
                try:
                    async with _httpx.AsyncClient(timeout=15.0, follow_redirects=True) as http:
                        spec_resp = await http.get(spec_url)
                        spec_resp.raise_for_status()
                        pages.append((f"{request.api_name} - OpenAPI Spec", spec_resp.text))
                except Exception:
                    pass

            # Crawl depth-2
            crawled_text, pages_crawled, crawled_pages = await _crawl_api_docs(
                base_url=request.api_docs_url,
                initial_html=initial_html,
                max_pages=50,
                max_depth=2,
            )
            # Add the initial page
            pages.append((f"{request.api_name} - Main Documentation", _extract_text_from_html(initial_html)))
            # Add all crawled pages
            for url, text in crawled_pages:
                # Use the URL path as a title
                from urllib.parse import urlparse
                path = urlparse(url).path.rstrip("/").split("/")[-1] or "index"
                pages.append((f"{request.api_name} - {path}", text))

            logger.info("mcp_fetch_docs_crawled", pages=len(pages), total_chars=sum(len(t) for _, t in pages))

            # Also search for OpenAPI spec on GitHub if we didn't find one
            if not any("OpenAPI" in title for title, _ in pages):
                try:
                    from app.services.mcp_builder_service import MCPBuilderService
                    svc = MCPBuilderService()
                    llm_settings = await get_active_llm_settings(session, supabase_jwt=_extract_supabase_jwt(raw_request))
                    spec_url = await svc._search_openapi_spec(request.api_name, llm_settings)
                    if spec_url:
                        async with _httpx.AsyncClient(timeout=15.0, follow_redirects=True) as http:
                            spec_resp = await http.get(spec_url)
                            spec_resp.raise_for_status()
                            pages.append((f"{request.api_name} - OpenAPI Spec (GitHub)", spec_resp.text))
                            logger.info("mcp_fetch_docs_openapi_found", url=spec_url)
                except Exception:
                    pass

            # Also search for SDK auth source code
            sdk_links = _re.findall(r'href="(https?://github\.com/[^"]+)"', initial_html)
            for sdk_url in sdk_links[:2]:  # Max 2 SDK repos
                try:
                    # Convert github.com URL to API URL to get file tree
                    parts = sdk_url.rstrip("/").split("/")
                    if len(parts) >= 5:
                        owner, repo = parts[3], parts[4]
                        api_url = f"https://api.github.com/repos/{owner}/{repo}/git/trees/main?recursive=1"
                        async with _httpx.AsyncClient(timeout=10.0) as http:
                            tree_resp = await http.get(api_url)
                            if tree_resp.status_code == 200:
                                tree = tree_resp.json().get("tree", [])
                                auth_files = [f for f in tree if f.get("type") == "blob" and
                                    any(kw in f.get("path", "").lower() for kw in ["auth", "sign", "hmac", "credential", "signature"])]
                                # Fetch auth files (max 5)
                                for af in auth_files[:5]:
                                    raw_url = f"https://raw.githubusercontent.com/{owner}/{repo}/main/{af['path']}"
                                    async with _httpx.AsyncClient(timeout=10.0) as http2:
                                        file_resp = await http2.get(raw_url)
                                        if file_resp.status_code == 200:
                                            pages.append((f"SDK - {af['path']}", file_resp.text))
                except Exception:
                    pass

    if not pages:
        raise HTTPException(status_code=400, detail="No documentation content could be fetched.")

    # Store all pages in KB
    kb_product_id = await store_mcp_docs_kb(
        server_name=safe_name,
        api_name=request.api_name,
        pages=pages,
    )

    logger.info("mcp_fetch_docs_complete", kb_product_id=kb_product_id, pages_stored=len(pages))

    return {
        "status": "success",
        "kb_product_id": kb_product_id,
        "pages_stored": len(pages),
        "summary": f"Fetched and stored {len(pages)} pages of documentation",
    }


class MCPAnalyzeKBRequest(BaseModel):
    api_name: str
    kb_product_id: str


@router.post("/analyze-kb")
async def analyze_kb_endpoints(
    request: MCPAnalyzeKBRequest,
    raw_request: Request,
    session: AsyncSession = Depends(get_session)
):
    """Analyze KB contents to discover endpoints — same flow as internal products."""
    logger.info("mcp_builder_analyze_kb_start", api_name=request.api_name, kb_product_id=request.kb_product_id)

    llm_settings = await get_active_llm_settings(session, supabase_jwt=_extract_supabase_jwt(raw_request))
    api_key = llm_settings.get("api_key")
    if not api_key:
        detail = llm_settings.get("llm_unavailable_detail", "LLM API key is required.")
        raise HTTPException(status_code=400, detail=detail)

    # Load KB content (clear cache first to ensure fresh data)
    from app.rag.kb_store import KnowledgeBaseStore, _kb_cache
    if request.kb_product_id in _kb_cache:
        del _kb_cache[request.kb_product_id]
    kb_store = KnowledgeBaseStore()
    kb_data = await kb_store.load(request.kb_product_id)
    if not kb_data:
        raise HTTPException(status_code=404, detail="No KB found. Fetch docs first.")

    # Extract all stored text from KB
    all_text = ""
    files = kb_data.get("files", {})
    for fname, fdata in files.items():
        if isinstance(fdata, dict):
            # KB stores content under "text" key
            all_text += fdata.get("text", fdata.get("content", "")) + "\n\n"
        elif isinstance(fdata, str):
            all_text += fdata + "\n\n"

    if not all_text.strip():
        raise HTTPException(status_code=404, detail="KB is empty. Fetch docs first.")

    # Use the SAME discovery flow as internal products
    from app.services.mcp_builder_service import MCPBuilderService
    svc = MCPBuilderService()
    result = await svc.discover_from_text(
        api_name=request.api_name,
        api_docs_text=all_text,
        api_key=llm_settings.get("api_key", ""),
        model=llm_settings.get("model_name", "gpt-4o"),
        api_url=llm_settings.get("api_url"),
        default_headers=llm_settings.get("default_headers", {}),
        provider=llm_settings.get("provider", "openai"),
    )

    return {
        "status": "success",
        "base_url": result.base_url,
        "auth_type": result.auth_type,
        "auth_details": getattr(result, "auth_details", None),
        "endpoints": [ep.model_dump() for ep in result.endpoints],
        "kb_product_id": request.kb_product_id,
    }


@router.post("/generate")
async def generate_mcp_server(
    request: MCPGeneratorRequest,
    raw_request: Request,
    session: AsyncSession = Depends(get_session)
):
    logger.info("mcp_builder_start", product_name=request.product_name, dest=request.destination_folder)

    agent = AgentService()

    # Initialize progress tracking
    safe_name_for_progress = re.sub(r'[^a-zA-Z0-9]', '_', request.product_name).lower().strip('_')
    safe_name_for_progress = re.sub(r'_+', '_', safe_name_for_progress)
    _generation_progress[safe_name_for_progress] = {
        "status": "running",
        "phase": "Initializing...",
        "logs": [],
        "iteration": 0,
        "result": None,
        "server_name": safe_name_for_progress,
    }

    # Compute module name early — needed in task prompts below
    safe_product_name = re.sub(r'[^a-zA-Z0-9]', '_', request.product_name).lower().strip('_')
    safe_product_name = re.sub(r'_+', '_', safe_product_name)
    module_name = f"{safe_product_name}_mcp"
    
    # Build destination: {user_chosen_dir}/mcp-servers/{server_name}
    # This ensures MCP files are always inside an mcp-servers subdirectory
    raw_dest = os.path.expanduser(request.destination_folder)
    raw_dest_path = Path(raw_dest)
    # If the path already ends with mcp-servers/{name}, use as-is (e.g. ~/mcp-servers/gta)
    # Otherwise, append mcp-servers/{name} to the user's chosen directory
    if raw_dest_path.parent.name == "mcp-servers":
        dest = str(raw_dest_path)
    else:
        dest = str(raw_dest_path / "mcp-servers" / safe_product_name)

    # Auto-increment folder name if it already exists (like file downloads: name, name_1, name_2)
    if os.path.exists(dest):
        base_dest = dest.rstrip('/')
        # Strip existing _N suffix to find the true base
        base_match = re.match(r'^(.+?)(?:_(\d+))?$', base_dest)
        base_path = base_match.group(1) if base_match else base_dest
        counter = 1
        while os.path.exists(f"{base_path}_{counter}"):
            counter += 1
        dest = f"{base_path}_{counter}"
        # Update module name to match
        new_folder_name = os.path.basename(dest)
        safe_product_name = re.sub(r'[^a-zA-Z0-9]', '_', new_folder_name).lower().strip('_')
        safe_product_name = re.sub(r'_+', '_', safe_product_name)
        module_name = f"{safe_product_name}_mcp"
        logger.info("mcp_builder_auto_increment", original=request.destination_folder, new_dest=dest)

    is_external = bool(request.api_docs_url or request.api_docs_text)

    if is_external:
        task_prompt = (
            f"Please generate a complete python MCP server for {request.product_name}. "
            f"This is an EXTERNAL API — you do NOT have a knowledge base for it.\n\n"
            f"Base URL: {request.api_base_url or 'Discover from the docs'}\n"
            f"Auth Type: {request.auth_type or 'bearer'}\n\n"
            f"The server MUST precisely implement the capabilities specified in the attached selected tools array.\n\n"
            f"CRITICAL INSTRUCTION: You are running in an automated background pipeline. "
            f"DO NOT WAIT FOR USER CONFIRMATION. The user has ALREADY selected and approved these exact tools in the UI. "
            f"SKIP Phase 1 entirely — the API documentation has already been analyzed and the endpoint details "
            f"are provided in your system prompt under EXTERNAL API MODE. "
            f"DO NOT call web_fetch() or web_search() — all the information you need is already in your context. "
            f"Proceed IMMEDIATELY to Phase 2 and write the FULL, complete implementation code for every single tool. "
            f"Do not write empty stubs or 'pass'.\n\n"
            f"FILE WRITING INSTRUCTION: You MUST use the write_file native tool to create files on disk.\n"
            f"Write the following files to {dest}:\n"
            f"1. {dest}/pyproject.toml (with dependencies: mcp, httpx, pydantic)\n"
            f"2. {dest}/src/{module_name}/__init__.py (empty)\n"
            f"3. {dest}/src/{module_name}/server.py (YOUR COMPLETE MCP SERVER CODE)\n"
            f"4. {dest}/README.md\n\n"
            f"CRITICAL: The generated code MUST sanitize API tokens/keys from environment variables. "
            f"Always use `token = os.getenv('...').strip('<>').strip('\"').strip(\"' \")` to handle cases where users "
            f"accidentally paste tokens with literal brackets or quotes.\n\n"
            f"NEVER output code into your chat response. ALL code must be saved to files using write_file()."
        )
    else:
        task_prompt = (
            f"Please generate a complete python MCP server for {request.product_name}. "
            f"The server MUST precisely implement the capabilities specified in the attached selected tools array.\n\n"
            f"CRITICAL INSTRUCTION: You are running in an automated background pipeline. "
            f"You MUST perform ALL of Phase 1 (all 6 queries) to deeply understand the product's API — "
            f"exact parameter names, types, route prefixes, auth patterns, and error formats. "
            f"However, do NOT wait for user confirmation at the end of Phase 1 — the user has ALREADY "
            f"selected and approved these exact tools in the UI. After completing Phase 1, proceed directly to Phase 2 "
            f"and write the FULL, complete implementation code for every single tool.\n\n"
            f"FILE WRITING INSTRUCTION: You MUST use the write_file native tool to create files on disk.\n"
            f"Write the following files to {dest}:\n"
            f"1. {dest}/pyproject.toml (with dependencies: mcp, httpx, pydantic)\n"
            f"2. {dest}/src/{module_name}/__init__.py (empty)\n"
            f"3. {dest}/src/{module_name}/server.py (YOUR COMPLETE MCP SERVER CODE)\n"
            f"4. {dest}/README.md\n\n"
            f"CRITICAL: The generated code MUST sanitize API tokens/keys from environment variables. "
            f"Always use `token = os.getenv('...').strip('<>').strip('\"').strip(\"' \")` to handle cases where users "
            f"accidentally paste tokens with literal brackets or quotes.\n\n"
            f"NEVER output code into your chat response. ALL code must be saved to files using write_file()."
        )
    
    mcp_config = {
        "language": "python",
        "output_dir": request.destination_folder,
        "selected_endpoints": request.selected_endpoints
    }

    # Pass external API context to the prompt builder
    if is_external:
        mcp_config["api_docs_url"] = request.api_docs_url
        mcp_config["api_base_url"] = request.api_base_url
        mcp_config["auth_type"] = request.auth_type
        mcp_config["auth_details"] = request.auth_details
        
        # When KB exists, agent queries it directly — skip injecting docs text
        # This saves context window space and lets the agent get focused results
        if not request.kb_product_id:
            # No KB — inject docs text into prompt as before
            docs_text = request.api_docs_text
            if docs_text:
                sliced_spec = mcp_builder_service.slice_openapi_for_endpoints(docs_text, request.selected_endpoints)
                if sliced_spec:
                    mcp_config["api_docs_text"] = sliced_spec
                else:
                    mcp_config["api_docs_text"] = docs_text[:50000] + ("..." if len(docs_text) > 50000 else "")
            elif request.api_docs_url:
                try:
                    import httpx
                    async with httpx.AsyncClient(timeout=10.0) as http:
                        resp = await http.get(request.api_docs_url)
                        sliced_spec = mcp_builder_service.slice_openapi_for_endpoints(resp.text, request.selected_endpoints)
                        if sliced_spec:
                            mcp_config["api_docs_text"] = sliced_spec
                            mcp_config["api_docs_url"] = None
                except Exception:
                    pass

            # Fallback: build structured summary from selected endpoints
            if not mcp_config.get("api_docs_text") and request.selected_endpoints:
                ep_lines = []
                for ep in request.selected_endpoints:
                    line = f"{ep.get('method', 'GET')} {ep.get('path', '/')} — {ep.get('description', '')}"
                    params = ep.get("parameters", [])
                    if params:
                        param_strs = []
                        for p in params:
                            req = " (required)" if p.get("required") else ""
                            param_strs.append(f"  - {p.get('name')}: {p.get('type', 'string')}{req} [{p.get('location', 'body')}]")
                        line += "\n" + "\n".join(param_strs)
                    ep_lines.append(line)
                auth_section = f"Auth Type: {request.auth_type or 'bearer'}"
                if request.auth_details:
                    auth_section += f"\nAuth Details: {request.auth_details}"
                endpoint_summary = (
                    f"API Base URL: {request.api_base_url or 'unknown'}\n"
                    f"{auth_section}\n\n"
                    f"Endpoints:\n\n" + "\n\n".join(ep_lines)
                )
                mcp_config["api_docs_text"] = endpoint_summary
                mcp_config["api_docs_url"] = None

    def _sse(event: str, data: Any) -> str:
        return f"event: {event}\ndata: {json.dumps(data)}\n\n"

    def _extract_mcp_config(final_answer_text: str) -> dict:
        """Build MCP config by inspecting what the agent actually generated on disk.

        1. Read pyproject.toml to find the real module name
        2. Scan server.py for os.environ/os.getenv calls to discover real env var names
        3. Fall back to conventions only if files are missing
        """
        expanded_dest = os.path.expanduser(dest)

        # ── 1. Discover real module name from pyproject.toml ──
        actual_module = module_name  # fallback
        pyproject_path = os.path.join(expanded_dest, "pyproject.toml")
        if os.path.isfile(pyproject_path):
            try:
                with open(pyproject_path) as f:
                    content = f.read()
                m = re.search(r'=\s*"([a-zA-Z0-9_]+)\.server:main"', content)
                if m:
                    actual_module = m.group(1)
            except Exception:
                pass

        # ── 2. Find server.py — check both src/ layout and flat layout ──
        server_path = None
        for candidate in [
            os.path.join(expanded_dest, "src", actual_module, "server.py"),
            os.path.join(expanded_dest, actual_module, "server.py"),
            os.path.join(expanded_dest, "server.py"),
        ]:
            if os.path.isfile(candidate):
                server_path = candidate
                break

        # ── 3. Extract env var names from actual server.py ──
        env_vars: dict = {}
        if server_path:
            try:
                with open(server_path) as f:
                    server_src = f.read()
                env_matches = re.findall(
                    r'os\.(?:environ(?:\.get)?\s*\[\s*["\']|environ\.get\s*\(\s*["\']|getenv\s*\(\s*["\'])([A-Z][A-Z0-9_]*)',
                    server_src,
                )
                for var_name in dict.fromkeys(env_matches):
                    env_vars[var_name] = f"<your-{var_name.lower()}>"
            except Exception:
                pass

        # Fallback: if no env vars found but auth is needed, use convention
        if not env_vars:
            auth = getattr(request, 'auth_type', None) or ''
            if auth and auth.lower() not in ('none', '', 'no auth', 'no_auth'):
                env_vars = {f"{safe_product_name.upper()}_API_TOKEN": "<your-token>"}

        # Also check agent's JSON output for env vars we might have missed
        match = re.search(r'```json\n(.*?)\n```', final_answer_text, re.DOTALL)
        if match:
            try:
                parsed = json.loads(match.group(1))
                servers = parsed.get("mcpServers", parsed)
                for srv_config in servers.values():
                    if isinstance(srv_config, dict) and "env" in srv_config:
                        for k, v in srv_config["env"].items():
                            if k not in env_vars:
                                env_vars[k] = v
                        break
            except Exception:
                pass

        # ── 4. Determine command — prefer .venv if it exists, else uv run ──
        venv_python = os.path.join(expanded_dest, ".venv", "bin", "python")
        if os.path.isfile(venv_python):
            command = venv_python
            args = ["-m", f"{actual_module}.server"]
        else:
            command = "uv"
            args = ["run", "--directory", expanded_dest, "python", "-m", f"{actual_module}.server"]

        config = {
            safe_product_name: {
                "command": command,
                "args": args,
            }
        }
        if env_vars:
            config[safe_product_name]["env"] = env_vars
        return config

    def _build_quick_start() -> str:
        expanded_dest = os.path.expanduser(dest)

        # Discover actual module name (same logic as config)
        actual_module = module_name
        pyproject_path = os.path.join(expanded_dest, "pyproject.toml")
        if os.path.isfile(pyproject_path):
            try:
                with open(pyproject_path) as f:
                    content = f.read()
                m = re.search(r'=\s*"([a-zA-Z0-9_]+)\.server:main"', content)
                if m:
                    actual_module = m.group(1)
            except Exception:
                pass

        venv_python = os.path.join(expanded_dest, ".venv", "bin", "python")
        if os.path.isfile(venv_python):
            run_cmd = f"{venv_python} -m {actual_module}.server"
        else:
            run_cmd = f"uv run --directory {dest} python -m {actual_module}.server"

        return (
            f"# Step 1: Install dependencies (only needed once)\n"
            f"cd {dest}\n"
            f"rm -f uv.lock\n"
            f"uv sync\n"
            f"\n"
            f"# Step 2: Start the MCP server\n"
            f"{run_cmd}\n"
            f"\n"
            f"# Step 3 (optional): Test with MCP Inspector\n"
            f"npx @modelcontextprotocol/inspector {run_cmd}"
        )

    # For external APIs with a KB, use the KB product_id so the agent gets
    # knowledge_base tool access to search the crawled/uploaded docs.
    agent_product_id = request.kb_product_id if request.kb_product_id else request.product_id

    state = {"final_answer": ""}

    async def _run_agent(task_str: str, max_iter: int):
        """Run the agent and return (final_answer, error)."""
        try:
            async for event in agent.execute_task(
                product_id=agent_product_id,
                task=task_str,
                session=session,
                mcp_builder_config=mcp_config,
                max_iterations=max_iter
            ):
                yield event
                if event.startswith("event: agent_answer_chunk"):
                    lines = event.split("\n")
                    if len(lines) >= 2 and lines[1].startswith("data: "):
                        try:
                            # Safely extract the chunk
                            data_json = json.loads(lines[1].replace("data: ", ""))
                            if data_json.get("done") and "chunk" in data_json:
                                state["final_answer"] = data_json["chunk"]
                        except Exception:
                            pass
        except Exception as e:
            logger.error("mcp_builder_agent_error", error=str(e))
            yield _sse("agent_status", {"step": "error", "message": f"Agent error: {str(e)}"})

    def _rescue_extract_server_code(answer: str) -> bool:
        """Last resort: if the agent dumped server code into its response text,
        extract it and write it to disk ourselves. Returns True if rescued."""
        if not answer:
            return False

        server_path = os.path.expanduser(f"{dest}/src/{module_name}/server.py")
        init_path = os.path.expanduser(f"{dest}/src/{module_name}/__init__.py")
        pyproject_path = os.path.expanduser(f"{dest}/pyproject.toml")

        # Look for Python code blocks that look like server code
        # Handle both closed (```python...```) and unclosed (```python... EOF) blocks
        code_blocks = re.findall(r'```python\s*\n(.*?)```', answer, re.DOTALL)

        # Also try unclosed blocks — agent often dumps code that gets truncated
        unclosed = re.findall(r'```python\s*\n(.+)', answer, re.DOTALL)
        for block in unclosed:
            # Remove any trailing ``` if partially present
            block = block.rstrip('`').strip()
            if block and block not in code_blocks:
                code_blocks.append(block)

        # Find the largest code block that looks like server code
        # (contains imports + function definitions — doesn't have to use FastMCP)
        best_block = None
        best_score = 0
        for block in code_blocks:
            score = 0
            if "import " in block:
                score += 1
            if "def " in block:
                score += 1
            if "@mcp.tool" in block:
                score += 10
            if "FastMCP" in block or "fastmcp" in block:
                score += 10
            if "httpx" in block or "AsyncClient" in block:
                score += 2
            if "async def" in block:
                score += 2
            if len(block) > 500:
                score += 3
            if score > best_score:
                best_score = score
                best_block = block

        if best_block and best_score >= 3:
            os.makedirs(os.path.dirname(server_path), exist_ok=True)
            with open(server_path, "w") as f:
                f.write(best_block)
            if not os.path.exists(init_path):
                with open(init_path, "w") as f:
                    f.write("")
            logger.info("mcp_builder_rescue_extracted", path=server_path, code_len=len(best_block), score=best_score)

            # Also look for pyproject.toml in the code blocks
            for block in code_blocks:
                if "[build-system]" in block or "[project]" in block:
                    if not os.path.exists(pyproject_path):
                        with open(pyproject_path, "w") as f:
                            f.write(block)
                        logger.info("mcp_builder_rescue_pyproject", path=pyproject_path)
                    break

            return True

        # Also check if the entire answer looks like Python code (no markdown)
        stripped = answer.strip()
        if stripped.startswith(("import ", "from ")) and ("def " in stripped) and len(stripped) > 500:
            os.makedirs(os.path.dirname(server_path), exist_ok=True)
            with open(server_path, "w") as f:
                f.write(stripped)
            if not os.path.exists(init_path):
                with open(init_path, "w") as f:
                    f.write("")
            logger.info("mcp_builder_rescue_raw_code", path=server_path, code_len=len(stripped))
            return True

        return False

    async def _validate_server() -> tuple[str | None, str]:
        """Validate the generated server. Returns (error_or_None, severity).

        severity: "missing" = server.py doesn't exist at all (needs full rewrite)
                  "syntax"  = server.py exists but has syntax/import errors
                  "incomplete" = server.py exists but has no tool functions
                  "ok"      = all checks passed
        """
        import subprocess
        import sys
        server_path = os.path.expanduser(f"{dest}/src/{module_name}/server.py")

        if not os.path.exists(server_path):
            return f"Server file not found at {server_path}", "missing"

        # Check file is not empty/stub
        try:
            content = open(server_path).read()
            if len(content.strip()) < 100:
                return f"Server file at {server_path} is essentially empty ({len(content)} chars)", "missing"
            # Count @mcp.tool() decorators as a proxy for implemented tools
            tool_count = content.count("@mcp.tool")
            if tool_count == 0:
                return f"Server file exists but contains 0 @mcp.tool() functions. Expected {len(request.selected_endpoints)} tools.", "incomplete"
        except Exception as e:
            return str(e), "syntax"

        # Syntax check
        try:
            result = subprocess.run(
                [sys.executable, "-m", "py_compile", server_path],
                capture_output=True, text=True, timeout=10
            )
            if result.returncode != 0:
                return result.stderr.strip(), "syntax"
        except Exception as e:
            return str(e), "syntax"

        return None, "ok"

    def _update_progress(**kwargs):
        """Update the in-memory progress store for this generation."""
        prog = _generation_progress.get(safe_name_for_progress)
        if prog:
            prog.update(kwargs)

    def _track_sse_event(event: str):
        """Parse an SSE event and update the progress store."""
        try:
            if "event: agent_status" in event:
                data_match = re.search(r'data: (.+)', event)
                if data_match:
                    data = json.loads(data_match.group(1))
                    _update_progress(phase=data.get("message", data.get("step", "")))
                    prog = _generation_progress.get(safe_name_for_progress)
                    if prog:
                        prog["logs"] = (prog.get("logs") or [])[-29:] + [data.get("message", "")]
            elif "event: agent_code" in event:
                data_match = re.search(r'data: (.+)', event)
                if data_match:
                    data = json.loads(data_match.group(1))
                    _update_progress(iteration=data.get("iteration", 0))
            elif "event: tool_output" in event:
                data_match = re.search(r'data: (.+)', event)
                if data_match:
                    data = json.loads(data_match.group(1))
                    output = (data.get("output") or "")[:120].replace("\n", " ")
                    if output:
                        prog = _generation_progress.get(safe_name_for_progress)
                        if prog:
                            prog["logs"] = (prog.get("logs") or [])[-29:] + [output]
            elif "event: mcp_result" in event:
                data_match = re.search(r'data: (.+)', event)
                if data_match:
                    data = json.loads(data_match.group(1))
                    status = "completed" if data.get("status") in ("success", "warning") else "failed"
                    _update_progress(status=status, result=data)
        except Exception:
            pass

    async def _stream_generation():
        # Ensure the Supabase JWT is available in ContextVar for agent's
        # internal get_active_llm_settings calls (StreamingResponse runs
        # in a wrapped context, but set it explicitly for safety).
        from app.services.gateway_session import gateway_supabase_token as _sb_tok
        _jwt = _extract_supabase_jwt(raw_request)
        _sb_tok_reset = _sb_tok.set(_jwt) if _jwt else None

        state["final_answer"] = ""
        _update_progress(phase="Agent starting...", status="running")

        # ── Attempt 1: Full generation ────────────────────────────
        async for event in _run_agent(task_prompt, max_iter=60):
            _track_sse_event(event)
            yield event

        final_answer = state["final_answer"]

        # ── Try rescue extraction before validation ─────────────
        # If agent dumped code into its response, save it to disk
        _rescue_extract_server_code(final_answer)

        # ── Server-side validation ────────────────────────────────
        validation_error, severity = await _validate_server()

        if validation_error:
            logger.warning("mcp_builder_validation_failed", error=validation_error, severity=severity)
            yield _sse("agent_status", {
                "step": "retry",
                "message": f"Validation failed ({severity}), auto-retrying: {validation_error[:150]}"
            })

            # Build retry prompt based on severity
            if severity == "missing":
                # Server.py doesn't exist — agent needs to write everything
                endpoint_names = ", ".join(
                    ep.get("suggested_tool_name", ep.get("path", "?"))
                    for ep in request.selected_endpoints[:20]
                )
                retry_prompt = (
                    f"CRITICAL: The MCP server generation FAILED. The file {dest}/src/{module_name}/server.py "
                    f"does NOT exist. You MUST write it now.\n\n"
                    f"Create the COMPLETE server at {dest}/src/{module_name}/server.py with:\n"
                    f"1. `from fastmcp import FastMCP` and `mcp = FastMCP(\"{safe_product_name}_mcp\")`\n"
                    f"2. One @mcp.tool() function for EACH of these tools: {endpoint_names}\n"
                    f"3. Each tool must make a real HTTP request using httpx.AsyncClient\n"
                    f"4. Also ensure {dest}/src/{module_name}/__init__.py exists\n"
                    f"5. Run `cd {dest} && uv sync` after writing the files\n\n"
                    f"IMPORTANT: To write the file, use the write_file native tool like this:\n"
                    f"write_file(\"{dest}/src/{module_name}/server.py\", '''<your complete server code here>''')\n\n"
                    f"Alternatively, you can use terminal to write:\n"
                    f"```python\n"
                    f"terminal(\"mkdir -p {dest}/src/{module_name} && cat > {dest}/src/{module_name}/server.py << 'PYEOF'\\n<your code>\\nPYEOF\")\n"
                    f"```\n\n"
                    f"DO NOT output code into your response text. You MUST use write_file() or terminal() to save files to disk."
                )
            elif severity == "incomplete":
                retry_prompt = (
                    f"The server file at {dest}/src/{module_name}/server.py exists but contains "
                    f"0 @mcp.tool() functions. It needs {len(request.selected_endpoints)} tools.\n\n"
                    f"Error: {validation_error}\n\n"
                    f"Read the current file, then add ALL missing @mcp.tool() functions. "
                    f"Each tool must make a real HTTP request. Do NOT write empty stubs."
                )
            else:
                # syntax error
                retry_prompt = (
                    f"The MCP server at {dest} has a syntax/import error:\n\n"
                    f"```\n{validation_error}\n```\n\n"
                    f"Read the file, fix the error with str_replace, then verify with "
                    f"`py_compile` and an import check."
                )

            _update_progress(phase="Retrying generation...")
            retry_agent = AgentService()
            async for event in retry_agent.execute_task(
                product_id=agent_product_id,
                task=retry_prompt,
                session=session,
                mcp_builder_config=mcp_config,
                max_iterations=20
            ):
                _track_sse_event(event)
                yield event
                if event.startswith("event: agent_answer_chunk"):
                    lines = event.split("\n")
                    if len(lines) >= 2 and lines[1].startswith("data: "):
                        try:
                            data_json = json.loads(lines[1].replace("data: ", ""))
                            if data_json.get("done") and "chunk" in data_json:
                                retry_answer = data_json["chunk"]
                                if retry_answer:
                                    final_answer = retry_answer
                        except Exception:
                            pass

            # Re-validate after retry
            validation_error_2, severity_2 = await _validate_server()
            if validation_error_2:
                logger.error("mcp_builder_retry_also_failed", error=validation_error_2, severity=severity_2)

                # Last resort: try to extract server code from the agent's response
                if severity_2 in ("missing", "incomplete"):
                    rescued = _rescue_extract_server_code(final_answer)
                    if rescued:
                        yield _sse("agent_status", {
                            "step": "rescue",
                            "message": "Extracted server code from agent response and saved to disk."
                        })

        # ── Final validation status ───────────────────────────────
        final_error, final_severity = await _validate_server()

        if final_error and final_severity == "missing":
            # Total failure — server.py never got created
            logger.error("mcp_builder_total_failure", error=final_error)
            fail_data = {
                "status": "error",
                "message": f"Generation failed: {final_error}. The LLM agent did not produce the server code. "
                           f"Please check your LLM API configuration and try again.",
                "mcpServers": None,
                "quickStartCommands": None,
            }
            _update_progress(status="failed", result=fail_data, phase="Generation failed")
            yield _sse("mcp_result", fail_data)
            return

        # ── Auto-setup: run uv sync if .venv is missing ──────────
        expanded_dest = os.path.expanduser(dest)
        venv_path = os.path.join(expanded_dest, ".venv", "bin", "python")
        if not os.path.exists(venv_path):
            import subprocess as _sp
            _update_progress(phase="Installing dependencies...")
            yield _sse("agent_status", {"step": "setup", "message": "Running uv sync to create virtual environment..."})
            try:
                setup_result = _sp.run(
                    ["uv", "sync"],
                    cwd=expanded_dest,
                    capture_output=True, text=True, timeout=120
                )
                if setup_result.returncode == 0:
                    logger.info("mcp_builder_auto_uv_sync", dest=expanded_dest, status="success")
                    yield _sse("tool_output", {"output": "uv sync completed — .venv created"})
                else:
                    logger.warning("mcp_builder_auto_uv_sync_failed", stderr=setup_result.stderr[:500])
                    yield _sse("tool_output", {"output": f"uv sync failed: {setup_result.stderr[:200]}"})
            except Exception as e:
                logger.warning("mcp_builder_auto_uv_sync_error", error=str(e))

        # ── Emit final result ─────────────────────────────────────
        mcp_servers_config = _extract_mcp_config(final_answer)
        status = "success" if not final_error else "warning"
        message = "Generation complete." if not final_error else f"Generation complete with warnings: {final_error}"

        logger.info("mcp_builder_complete", config=mcp_servers_config, status=status)

        # ── Save to registry ─────────────────────────────────────
        try:
            from app.models.mcp_server import McpServer
            source_type = "internal"
            if is_external:
                if request.api_docs_url:
                    source_type = "external_url"
                elif request.api_docs_text:
                    source_type = "external_text"
                else:
                    source_type = "external_upload"

            new_server = McpServer(
                name=safe_product_name,
                product_name=request.product_name,
                destination_folder=dest,
                module_name=module_name,
                mcp_config_json=mcp_servers_config,
                quick_start_commands=_build_quick_start(),
                selected_endpoints_json=request.selected_endpoints,
                source_type=source_type,
                api_docs_url=request.api_docs_url,
                api_base_url=request.api_base_url,
                auth_type=request.auth_type,
                kb_product_id=request.kb_product_id,
            )
            session.add(new_server)
            await session.commit()
            logger.info("mcp_server_saved", server_id=new_server.server_id, name=safe_product_name)
        except Exception as exc:
            logger.warning("mcp_server_save_failed", error=str(exc))

        success_data = {
            "status": status,
            "mcpServers": mcp_servers_config,
            "quickStartCommands": _build_quick_start(),
            "message": message,
        }
        _update_progress(
            status="completed" if status == "success" else "completed",
            result=success_data,
            phase="Generation complete"
        )
        yield _sse("mcp_result", success_data)

    return StreamingResponse(
        _stream_generation(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )
